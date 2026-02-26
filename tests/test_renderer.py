import re
from pathlib import Path

from docx import Document
from lxml import etree

from policy_localiser.engine.models import ProcessingStatus
from policy_localiser.engine.renderer import PolicyRenderer


class TestPolicyRenderer:
    def test_replaces_text_placeholders_in_body(self, template_path, logos_dir, stm_school, tmp_path):
        renderer = PolicyRenderer()
        output = tmp_path / "output.docx"

        result = renderer.render(
            template_path=template_path,
            logo_path=logos_dir / "STM.png",
            school=stm_school,
            output_path=output,
            run_id="test001",
        )

        assert result.status == ProcessingStatus.SUCCESS
        assert output.exists()

        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "St Mary's Primary School" in full_text
        assert "(07) 3000 1234" in full_text
        assert "admin@stmarys.qld.edu.au" in full_text

    def test_no_unreplaced_placeholders(self, template_path, logos_dir, stm_school, tmp_path):
        renderer = PolicyRenderer()
        output = tmp_path / "output.docx"

        renderer.render(
            template_path=template_path,
            logo_path=logos_dir / "STM.png",
            school=stm_school,
            output_path=output,
            run_id="test002",
        )

        doc = Document(str(output))
        full_text = " ".join(p.text for p in doc.paragraphs)
        remaining = re.findall(r"\{\{.*?\}\}", full_text)
        assert remaining == [], f"Unreplaced placeholders: {remaining}"

    def test_replaces_table_data(self, template_path, logos_dir, stm_school, tmp_path):
        renderer = PolicyRenderer()
        output = tmp_path / "output.docx"

        renderer.render(
            template_path=template_path,
            logo_path=logos_dir / "STM.png",
            school=stm_school,
            output_path=output,
            run_id="test003",
        )

        doc = Document(str(output))
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "

        assert "STM" in table_text
        assert "St Mary's Parish" in table_text
        assert "12 345 678 901" in table_text

    def test_header_image_present(self, template_path, logos_dir, stm_school, tmp_path):
        renderer = PolicyRenderer()
        output = tmp_path / "output.docx"

        renderer.render(
            template_path=template_path,
            logo_path=logos_dir / "STM.png",
            school=stm_school,
            output_path=output,
            run_id="test004",
        )

        doc = Document(str(output))
        for section in doc.sections:
            header_xml = etree.tostring(section.header._element, encoding="unicode")
            assert "blip" in header_xml, "No image found in header"

    def test_footer_field_preserved(self, template_path, logos_dir, stm_school, tmp_path):
        renderer = PolicyRenderer()
        output = tmp_path / "output.docx"

        renderer.render(
            template_path=template_path,
            logo_path=logos_dir / "STM.png",
            school=stm_school,
            output_path=output,
            run_id="test005",
        )

        doc = Document(str(output))
        for section in doc.sections:
            footer_xml = etree.tostring(section.footer._element, encoding="unicode")
            assert "fldChar" in footer_xml or "PAGE" in footer_xml, \
                "Footer page field was destroyed"

    def test_error_on_missing_logo(self, template_path, stm_school, tmp_path):
        renderer = PolicyRenderer()
        output = tmp_path / "output.docx"

        result = renderer.render(
            template_path=template_path,
            logo_path=Path("/nonexistent/logo.png"),
            school=stm_school,
            output_path=output,
            run_id="test006",
        )

        assert result.status == ProcessingStatus.ERROR
        assert result.error_message is not None

    def test_different_schools_produce_different_output(
        self, template_path, logos_dir, sample_schools, tmp_path
    ):
        renderer = PolicyRenderer()
        outputs = {}
        for school in sample_schools[:2]:
            output = tmp_path / f"{school.SchoolCode}.docx"
            renderer.render(
                template_path=template_path,
                logo_path=logos_dir / f"{school.SchoolCode}.png",
                school=school,
                output_path=output,
                run_id="test007",
            )
            doc = Document(str(output))
            outputs[school.SchoolCode] = "\n".join(p.text for p in doc.paragraphs)

        # STM and HFC should have different content
        assert "St Mary's Primary School" in outputs["STM"]
        assert "Holy Family College" in outputs["HFC"]
        assert "St Mary's Primary School" not in outputs["HFC"]
