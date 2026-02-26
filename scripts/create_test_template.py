"""Generate a sample policy template .docx with placeholders for testing.

Creates:
  - tests/fixtures/templates/Sample_Policy.docx
  - tests/fixtures/logos/STM.png
  - tests/fixtures/logos/HFC.png
  - tests/fixtures/logos/logo_placeholder.png  (used inside the template header)
"""

import sys
from pathlib import Path

# Add src to path
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont


FIXTURES_DIR = Path(__file__).parent.parent / "tests" / "fixtures"


def create_placeholder_logo(output_path: Path, text: str = "LOGO", color: str = "#CCCCCC"):
    """Create a simple placeholder PNG image."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (200, 80), color)
    draw = ImageDraw.Draw(img)
    # Draw text centered
    try:
        font = ImageFont.truetype("arial.ttf", 24)
    except OSError:
        font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), text, font=font)
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]
    x = (200 - text_width) // 2
    y = (80 - text_height) // 2
    draw.text((x, y), text, fill="white", font=font)
    img.save(str(output_path))
    print(f"  Created: {output_path}")


def create_school_logo(output_path: Path, school_code: str, color: str):
    """Create a coloured logo PNG for a test school."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (200, 80), color)
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 20)
    except OSError:
        font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), school_code, font=font)
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]
    x = (200 - text_width) // 2
    y = (80 - text_height) // 2
    draw.text((x, y), school_code, fill="white", font=font)
    img.save(str(output_path))
    print(f"  Created: {output_path}")


def create_sample_template(output_path: Path, logo_placeholder_path: Path):
    """Create a sample .docx policy template with placeholders."""
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # --- Header with placeholder logo (top-right) ---
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Create a table in the header for layout: left cell = title, right cell = logo
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left cell: policy title
    left_cell = header_table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = left_para.add_run("{{ShortName}} — Policy Document")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Right cell: placeholder logo image
    right_cell = header_table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.add_run().add_picture(str(logo_placeholder_path), width=Cm(3.5))

    # Remove borders from header table
    from docx.oxml.ns import qn
    tbl = header_table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)

    # --- Footer with page number field ---
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field code (complex field)
    from docx.oxml import OxmlElement
    run_elem = OxmlElement("w:r")
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run_elem.append(fld_char_begin)
    footer_para._element.append(run_elem)

    run_elem2 = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    run_elem2.append(instr_text)
    footer_para._element.append(run_elem2)

    run_elem3 = OxmlElement("w:r")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_elem3.append(fld_char_end)
    footer_para._element.append(run_elem3)

    # --- Body: Title ---
    title = doc.add_heading("Enrolment Policy", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Body: School info paragraph ---
    doc.add_heading("1. School Information", level=2)
    p = doc.add_paragraph()
    p.add_run("This policy applies to ")
    p.add_run("{{Title}}").bold = True
    p.add_run(" (\"{{ShortName}}\"), located at {{SchoolAddress}}, {{Suburb}}, {{State}} {{PostCode}}.")

    doc.add_paragraph(
        "The school is classified as a {{SchoolType}} school within the "
        "{{DiocesanRegion}} region of the diocese."
    )

    # --- Body: Contact details ---
    doc.add_heading("2. Contact Details", level=2)
    p2 = doc.add_paragraph()
    p2.add_run("{{PrincipalTitle}}: ").bold = True
    p2.add_run("{{PrincipalName}}")

    p3 = doc.add_paragraph()
    p3.add_run("Phone: ").bold = True
    p3.add_run("{{SchoolPhone}}")

    p4 = doc.add_paragraph()
    p4.add_run("Email: ").bold = True
    p4.add_run("{{SchoolEmail}}")

    p5 = doc.add_paragraph()
    p5.add_run("Website: ").bold = True
    p5.add_run("{{SchoolWebsite}}")

    # --- Body: Table with placeholders ---
    doc.add_heading("3. Key Details", level=2)
    table = doc.add_table(rows=4, cols=2, style="Table Grid")
    table.cell(0, 0).text = "School Code"
    table.cell(0, 1).text = "{{SchoolCode}}"
    table.cell(1, 0).text = "Parish"
    table.cell(1, 1).text = "{{Parish}}"
    table.cell(2, 0).text = "ABN"
    table.cell(2, 1).text = "{{ABN}}"
    table.cell(3, 0).text = "Established"
    table.cell(3, 1).text = "{{EstablishedYear}}"

    # --- Body: Policy content ---
    doc.add_heading("4. Policy Statement", level=2)
    doc.add_paragraph(
        "{{ShortName}} is committed to providing a safe, supportive, and inclusive "
        "learning environment for all students. This policy outlines the school's "
        "approach to enrolment and the procedures that families must follow."
    )

    doc.add_paragraph(
        "For further information, please contact {{PrincipalName}} at "
        "{{SchoolEmail}} or phone {{SchoolPhone}}."
    )

    doc.save(str(output_path))
    print(f"  Created: {output_path}")


def main():
    print("Creating test fixtures...")

    # 1. Create placeholder logo (used inside the template)
    logo_placeholder = FIXTURES_DIR / "logos" / "logo_placeholder.png"
    create_placeholder_logo(logo_placeholder, "LOGO", "#999999")

    # 2. Create school logos
    create_school_logo(FIXTURES_DIR / "logos" / "STM.png", "STM", "#2E86AB")
    create_school_logo(FIXTURES_DIR / "logos" / "HFC.png", "HFC", "#A23B72")
    create_school_logo(FIXTURES_DIR / "logos" / "SJV.png", "SJV", "#2CA58D")

    # 3. Create sample template
    create_sample_template(
        FIXTURES_DIR / "templates" / "Sample_Policy.docx",
        logo_placeholder,
    )

    print("\nAll fixtures created successfully!")


if __name__ == "__main__":
    main()
